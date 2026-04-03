#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
舆情监测日报文档读取脚本

功能：读取 .docx 格式的日报文档并提取纯文本内容
参数：file_path - 文档路径
输出：文档中的所有文本内容
"""

import sys
from docx import Document


def read_docx(file_path: str) -> str:
    """
    读取 .docx 文档并提取纯文本内容
    
    Args:
        file_path: 文档路径（相对路径或绝对路径）
    
    Returns:
        str: 文档中的所有文本内容
    
    Raises:
        FileNotFoundError: 文档不存在
        ValueError: 文档格式不正确或无法读取
    """
    try:
        # 加载文档
        doc = Document(file_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 过滤空段落
                paragraphs.append(text)
        
        # 提取表格中的文本（如果有）
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(' | '.join(row_text))
        
        # 合并所有文本
        full_text = '\n'.join(paragraphs)
        
        if not full_text:
            raise ValueError("文档内容为空或无法读取")
        
        return full_text
        
    except FileNotFoundError:
        raise FileNotFoundError(f"文档不存在: {file_path}")
    except Exception as e:
        raise ValueError(f"读取文档失败: {str(e)}")


def main():
    """主函数：处理命令行参数并执行文档读取"""
    if len(sys.argv) != 2:
        print("使用方法: python read_docx.py <文档路径>")
        print("示例: python read_docx.py ./daily_report.docx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        text = read_docx(file_path)
        print(text)
    except (FileNotFoundError, ValueError) as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
