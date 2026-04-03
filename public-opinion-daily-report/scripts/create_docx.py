#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成舆情监测日报 Word 文档

功能：根据模板生成两个 Word 文档
1. [日期]日报概述.docx - 包含标题、监测时间段、传播概述、舆情数据
2. [日期]高频正面信息.docx - 包含高频正面信息和高管高频正面信息
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_daily_overview(date_str, content):
    """
    创建日报概述文档
    
    Args:
        date_str: 日期字符串，格式如 "0327"
        content: 文档内容字典，包含以下字段：
            - title: 标题
            - time_range: 监测时间段
            - overview: 传播概述
            - data: 舆情数据（包含3个子项）
    """
    doc = Document()
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run(content['title'])
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # 添加监测时间段
    time_para = doc.add_paragraph()
    time_para.add_run(content['time_range'])
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加传播概述
    overview_para = doc.add_paragraph()
    overview_para.add_run(content['overview'])
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加舆情数据标题
    data_title = doc.add_paragraph()
    data_title_run = data_title.add_run('【舆情数据】')
    data_title_run.font.bold = True
    
    # 添加舆情数据内容
    for item in content['data']:
        doc.add_paragraph(item)
    
    # 保存文档
    filename = f"{date_str}日报概述.docx"
    doc.save(filename)
    return filename


def create_positive_info(date_str, content):
    """
    创建高频正面信息文档
    
    Args:
        date_str: 日期字符串，格式如 "0327"
        content: 文档内容字典，包含以下字段：
            - positive_info: 高频正面信息列表，每项包含标题和链接
    """
    doc = Document()
    
    # 添加高频正面信息标题
    title = doc.add_paragraph()
    title_run = title.add_run('【高频正面信息】')
    title_run.font.bold = True
    
    # 添加高频正面信息内容
    for idx, item in enumerate(content['positive_info'], 1):
        para = doc.add_paragraph()
        para.add_run(f"{idx}⃣{item['title']}")
        doc.add_paragraph(item['url'])
    
    # 保存文档
    filename = f"{date_str}高频正面信息.docx"
    doc.save(filename)
    return filename


def main():
    """主函数：示例用法"""
    print("使用方法：")
    print("1. create_daily_overview(date_str, content) - 创建日报概述文档")
    print("2. create_positive_info(date_str, content) - 创建高频正面信息文档")
    print("\n此脚本由智能体调用，直接传入参数生成文档")


if __name__ == "__main__":
    main()
